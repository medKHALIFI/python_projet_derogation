# -*- coding: utf-8 -*-
"""
/***************************************************************************
 miniprojet
                                 A QGIS plugin
 miniprojet
                              -------------------
        begin                : 2018-05-21
        git sha              : $Format:%H$
        copyright            : (C) 2018 by hanaa khoj et khalifi mohamed
        email                : hanaa.khoj@gmail.com
 ***************************************************************************/

/***************************************************************************
 *                                                                         *
 *   This program is free software; you can redistribute it and/or modify  *
 *   it under the terms of the GNU General Public License as published by  *
 *   the Free Software Foundation; either version 2 of the License, or     *
 *   (at your option) any later version.                                   *
 *                                                                         *
 ***************************************************************************/
"""


# Initialize Qt resources from file resources.py
import resources
# Import the code for the dialog
from miniprojet_dialog import miniprojetDialog
import os.path
from PyQt4.QtCore import *
from PyQt4.QtGui import *
from qgis.core import *
from qgis import *
from qgis.utils import iface
import sys
reload(sys)
sys.setdefaultencoding("utf-8")

from qgis.gui import *
from qgis.core import QgsApplication
from docx import Document
import docx 




class miniprojet:
    """QGIS Plugin Implementation."""

    def __init__(self, iface):
        """Constructor.

        :param iface: An interface instance that will be passed to this class
            which provides the hook by which you can manipulate the QGIS
            application at run time.
        :type iface: QgisInterface
        """
        # Save reference to the QGIS interface
        self.iface = iface
        # initialize plugin directory
        self.plugin_dir = os.path.dirname(__file__)
        # initialize locale
        locale = QSettings().value('locale/userLocale')[0:2]
        locale_path = os.path.join(
            self.plugin_dir,
            'i18n',
            'miniprojet_{}.qm'.format(locale))

        if os.path.exists(locale_path):
            self.translator = QTranslator()
            self.translator.load(locale_path)

            if qVersion() > '4.3.3':
                QCoreApplication.installTranslator(self.translator)


        # Declare instance attributes
        self.actions = []
        self.menu = self.tr(u'&miniprojet')
        # TODO: We are going to let the user set this up in a future iteration
        self.toolbar = self.iface.addToolBar(u'miniprojet')
        self.toolbar.setObjectName(u'miniprojet')

    # noinspection PyMethodMayBeStatic
    def tr(self, message):
        """Get the translation for a string using Qt translation API.

        We implement this ourselves since we do not inherit QObject.

        :param message: String for translation.
        :type message: str, QString

        :returns: Translated version of message.
        :rtype: QString
        """
        # noinspection PyTypeChecker,PyArgumentList,PyCallByClass
        return QCoreApplication.translate('miniprojet', message)


    def add_action(
        self,
        icon_path,
        text,
        callback,
        enabled_flag=True,
        add_to_menu=True,
        add_to_toolbar=True,
        status_tip=None,
        whats_this=None,
        parent=None):
        """Add a toolbar icon to the toolbar.

        :param icon_path: Path to the icon for this action. Can be a resource
            path (e.g. ':/plugins/foo/bar.png') or a normal file system path.
        :type icon_path: str

        :param text: Text that should be shown in menu items for this action.
        :type text: str

        :param callback: Function to be called when the action is triggered.
        :type callback: function

        :param enabled_flag: A flag indicating if the action should be enabled
            by default. Defaults to True.
        :type enabled_flag: bool

        :param add_to_menu: Flag indicating whether the action should also
            be added to the menu. Defaults to True.
        :type add_to_menu: bool

        :param add_to_toolbar: Flag indicating whether the action should also
            be added to the toolbar. Defaults to True.
        :type add_to_toolbar: bool

        :param status_tip: Optional text to show in a popup when mouse pointer
            hovers over the action.
        :type status_tip: str

        :param parent: Parent widget for the new action. Defaults None.
        :type parent: QWidget

        :param whats_this: Optional text to show in the status bar when the
            mouse pointer hovers over the action.

        :returns: The action that was created. Note that the action is also
            added to self.actions list.
        :rtype: QAction
        """

        # Create the dialog (after translation) and keep reference
        self.dlg = miniprojetDialog()

        self.dlg.pushButtonBuffer.clicked.connect(self.buffer)
        self.dlg.pushButton_create.clicked.connect(self.createProject)
        self.dlg.pushButtongo.clicked.connect(self.select_layer)
        self.dlg.comboBoxLayers.currentIndexChanged.connect(self.AfficherIntersection)
        self.dlg.tableWidget.currentItemChanged.connect(self.Zoom)
        self.dlg.pushbuttom_refresh1.clicked.connect(self.refresh1)
        self.dlg.pushbuttom_refresh2.clicked.connect(self.refresh2)
        self.dlg.pushbuttom_assiettes.clicked.connect(self.AfficherAssiettes)
        self.dlg.pushButton_carte.clicked.connect(self.exportMapToPDF)
        self.dlg.pushButton_2.clicked.connect(self.exportToWord)
        self.dlg.tableWidget_2.currentItemChanged.connect(self.Zoom1)

        icon = QIcon(icon_path)
        action = QAction(icon, text, parent)
        action.triggered.connect(callback)
        action.setEnabled(enabled_flag)

        if status_tip is not None:
            action.setStatusTip(status_tip)

        if whats_this is not None:
            action.setWhatsThis(whats_this)

        if add_to_toolbar:
            self.toolbar.addAction(action)

        if add_to_menu:
            self.iface.addPluginToVectorMenu(
                self.menu,
                action)

        self.actions.append(action)

        return action

    def initGui(self):
        """Create the menu entries and toolbar icons inside the QGIS GUI."""

        icon_path = ':/plugins/miniprojet/icon.png'
        self.add_action(
            icon_path,
            text=self.tr(u'miniprojet'),
            callback=self.run,
            parent=self.iface.mainWindow())


    def unload(self):
        """Removes the plugin menu item and icon from QGIS GUI."""
        for action in self.actions:
            self.iface.removePluginVectorMenu(
                self.tr(u'&miniprojet'),
                action)
            self.iface.removeToolBarIcon(action)
        # remove the toolbar
        del self.toolbar


    def FindLayerByName(self,NameLayer):
        layer = None
        for lyr in QgsMapLayerRegistry.instance().mapLayers().values():
           if lyr.name() == NameLayer:
               layer = lyr
               break
        return layer

    def createProject(self):
        x=self.dlg.lineEditX.text()
        textx=x
        x=float(x)
        y=self.dlg.lineEditY.text()
        texty=y
        y=float(y)
        vl=self.FindLayerByName('couche_projet')
        # CRS
        crsType = QSettings().value('/Projections/defaultBehaviour') 
        QSettings().setValue('/Projections/defaultBehaviour','')
        target_crs = QgsCoordinateReferenceSystem()
        target_crs.createFromId(100000, QgsCoordinateReferenceSystem.InternalCrsId)
        if vl==None :
            vl = QgsVectorLayer("Point","couche_projet", "memory") 
            vl.setCrs(target_crs)
            print vl.crs().authid()
            QSettings().setValue('/Projections/defaultBehaviour',crsType)

        pr = vl.dataProvider() 
        vl.startEditing()
        pr.addAttributes([QgsField("X", QVariant.String),QgsField("Y", QVariant.Double)])

        vl.selectAll() #ca c'est pour supprimer les anciens buffers 
        ids = [f.id() for f in vl.selectedFeatures()]
        for fid in ids:
            vl.deleteFeature( fid )
        
        feat = QgsFeature()
        feat.setGeometry(QgsGeometry.fromPoint(QgsPoint(x,y)))
        feat.setAttributes([textx,texty])
        pr.addFeatures([feat])
        vl.commitChanges()
        vl.updateExtents()
        QgsMapLayerRegistry.instance().addMapLayer(vl)

        



    def select_attributes(self): #Sélectionner les attributs de la couche déja sélectionnée

        V_layers=self.iface.legendInterface().layers()
        selectedLayerIndex = self.dlg.layers.currentIndex()
        selectedLayer = V_layers[selectedLayerIndex]
        # Identify fields of the selected layer
        return selectedLayersel  
    

    def select_layer(self): #Selectionner les couches qui se trouve dans le canva.

        global layer_list
        self.dlg.comboBoxLayers.clear()
        layers=self.iface.legendInterface().layers()
        layer_list=[]
        for layer in layers :
            if layer.name() != "buffer" and layer.name() != "couche_projet" and layer.name()!="OpenStreetMap" :
                layer_list.append(layer.name())

        self.dlg.comboBoxLayers.addItems(layer_list)


    def Zoom(self):
        canvas = self.iface.mapCanvas()
        cLayer = self.FindLayerByName(self.dlg.comboBoxLayers.currentText())
        self.iface.setActiveLayer(cLayer)
        cLayer.removeSelection()
        indexes = self.dlg.tableWidget.selectionModel().selectedRows()
        ligneSelectionnee=0
        for index in sorted(indexes):
            ligneSelectionnee=index.row()
        print self.dlg.tableWidget.item(ligneSelectionnee,0).text()
        expression= "\"OBJECTID\"="+self.dlg.tableWidget.item(ligneSelectionnee,0).text()
        expr = QgsExpression( expression)
        it = cLayer.getFeatures( QgsFeatureRequest( expr ) )
        ids = [i.id() for i in it]
        cLayer.setSelectedFeatures( ids )
        canvas.zoomToSelected(cLayer)

    def Zoom1(self):
        canvas = self.iface.mapCanvas()
        cLayer = self.FindLayerByName(self.dlg.comboBoxLayers.currentText())
        self.iface.setActiveLayer(cLayer)
        cLayer.removeSelection()

        indexes = self.dlg.tableWidget_2.selectionModel().selectedRows()
        ligneSelectionnee=0
        for index in sorted(indexes):
            ligneSelectionnee=index.row()
        
        # self.dlg.tableWidget.item(ligneSelectionnee,0).text()
        expression= "\"OBJECTID\"="+self.dlg.tableWidget.item(ligneSelectionnee,0).text()
        expr = QgsExpression( expression)
        it = cLayer.getFeatures( QgsFeatureRequest( expr ) )
        ids = [i.id() for i in it]
        cLayer.setSelectedFeatures( ids )
        canvas.zoomToSelected(cLayer)

    def refresh1(self):
        self.dlg.lineEditX.clear()
        self.dlg.lineEditY.clear()

    def refresh2(self):
        self.dlg.spinBox.clear()
    

    def AfficherAssiettes(self):
        if self.h==1:
            

            canvas = self.iface.mapCanvas()
            areas = []
            
            line_layer = self.FindLayerByName('buffer')
            area_layer = self.FindLayerByName(self.dlg.comboBoxLayers.currentText())
            prov = area_layer.dataProvider()
            fields = prov.fields()
            self.dlg.tableWidget.setColumnCount(7)
            self.dlg.tableWidget.setRowCount(1)
            self.dlg.tableWidget_2.setRowCount(1)
            self.dlg.tableWidget_2.setColumnCount(1)
            i=0
            for field in fields:
                if field.name() == "OBJECTID" or field.name() == "REGIME_FON" or field.name() == "REFERENCE_" or field.name() == "CERCLE" or field.name() == "COMMUNE" or field.name() == "STATUT_FON" or field.name() == "SUPERFICIE" :
                    self.dlg.tableWidget.setHorizontalHeaderItem (i,QTableWidgetItem(field.name()))
                    i=i+1 

            
            i=0
            self.dlg.tableWidget_2.setHorizontalHeaderItem (i,QTableWidgetItem("Proportionn"))
            for line_feature in line_layer.getFeatures():
                i=0
                for area_feature in area_layer.getFeatures():
                   
                    if (area_feature.geometry().intersects(line_feature.geometry())==True) :
                        areas.append(area_feature.id())

                        area_feature.geometry().area()
                        #QMessageBox.information(None, "DEBUG:",str(area_feature.geometry().area()))
                        if area_feature.geometry().area() > 10000 :
                            self.dlg.tableWidget.insertRow(i+1)
                            self.dlg.tableWidget_2.insertRow(i+1)
                            #self.dlg.tableW.setRowCount(j)
                            geom = line_feature.geometry().intersection(area_feature.geometry())
                            percent=(float(geom.area())/float(area_feature.geometry().area()))*100
                            self.dlg.tableWidget_2.setItem(i,0,QTableWidgetItem(str(percent)))
                            j=0
                            k=0

                            for field in fields:
                                if field.name() == "OBJECTID" or field.name() == "REGIME_FON" or field.name() == "REFERENCE_" or field.name() == "CERCLE" or field.name() == "COMMUNE" or field.name() == "STATUT_FON" or field.name() == "SUPERFICIE" :
                                    s=''
                                    try :
                                        s=str(area_feature.attributes()[j])
                                        self.dlg.tableWidget.setItem(i,k,QTableWidgetItem(s))
                                    except :
                                        pass
                                    k=k+1
                                j=j+1
                            i=i+1
                        
                self.dlg.tableWidget.removeRow (i)
                self.dlg.tableWidget_2.removeRow (i)
            legend = self.iface.legendInterface() 
            #for layer in canvas.layers():
                #if layer.type() == layer.VectorLayer and layer.name() != "OpenStreetMap" :
                    #layer.removeSelection()
                     # access the legend
                    #legend.setLayerVisible(layer, False) 
            canvas.refresh()
            legend.setLayerVisible(area_layer, True)
            legend.setLayerVisible(line_layer, True)
            project=self.FindLayerByName('couche_projet')
            legend.setLayerVisible(project, True)
            area_layer.select(areas)
            canvas.zoomToSelected(area_layer)
            self.dlg.label_3.setText(" ")
            self.dlg.label_3.setText("LE NOMBRE DES ASSIETTES FONCIERES QUI DEPASSENT 1ha DE SURFACE EST: " + str(self.dlg.tableWidget.rowCount()) + " !")

    def buffer(self):
            self.h=1
            canvas = self.iface.mapCanvas()
            x=self.dlg.lineEditX.text()
            textx=x
            y=self.dlg.lineEditY.text()
            texty=y
            layerXY=self.FindLayerByName('couche_projet')
            layer_buff = self.FindLayerByName('buffer')

            crsType = QSettings().value('/Projections/defaultBehaviour')
            QSettings().setValue('/Projections/defaultBehaviour','')
            target_crs = QgsCoordinateReferenceSystem()
            target_crs.createFromId(100000, QgsCoordinateReferenceSystem.InternalCrsId)
            if layer_buff==None :
                layer_buff =  QgsVectorLayer("Polygon", "buffer" , "memory")
                layer_buff.setCrs(target_crs)
                print layer_buff.crs().authid()
                QSettings().setValue('/Projections/defaultBehaviour',crsType) 


            pr = layer_buff.dataProvider() 
            layer_buff.startEditing()
            pr.addAttributes([QgsField("X_Centre", QVariant.Double),QgsField("Y_Centre", QVariant.Double),QgsField("Perimetre", QVariant.Double)])
            bfr=self.dlg.spinBox.value()
            bfr=float(bfr)
            
            layer_buff.selectAll()
            #ca c'est pour supprimer les anciens buffers 
            ids = [f.id() for f in layer_buff.selectedFeatures()]
            layer_buff.startEditing()
            for fid in ids:
               layer_buff.deleteFeature( fid )
            layer_buff.commitChanges()
            
            
            for elem in layerXY.getFeatures():
                ensembAttribut = elem.attributes()
                geom = elem.geometry()
                buffer = geom.buffer(bfr,20)
                seg = QgsFeature()
                seg.setGeometry(buffer)
                seg.setAttributes([textx,texty,bfr])
                pr.addFeatures([seg])
                layer_buff.updateExtents()
            QgsMapLayerRegistry.instance().addMapLayers([layer_buff])
            layer_buff.isValid()


    def run(self):
        """Run method that performs all the real work"""
        # show the dialog
        self.dlg.show()
        # Run the dialog event loop
        result = self.dlg.exec_()
        # See if OK was pressed
        if result:
            # Do something useful here - delete the line containing pass and
            # substitute with your code.
            pass

    def exportMapToPDF(self):  

        mapRenderer = iface.mapCanvas().mapRenderer()
        c = QgsComposition(mapRenderer)
        c.setPlotStyle(QgsComposition.Print)
        x, y = 0,0
        w, h = c.paperWidth(), c.paperHeight()
        composerMap = QgsComposerMap(c, x ,y, w, h)
        c.addItem(composerMap)
        
        #legende
        legend = QgsComposerLegend(c)
        legend.model().setLayerSet(mapRenderer.layerSet())
        legend.setItemPosition(0,150)
        c.addItem(legend)
        #title
        composerLabel = QgsComposerLabel(c)
        composerLabel.setText("Les projets ayant une derogation")
        composerLabel.setFont(QFont("Cambria", 15, QFont.Bold))
        composerLabel.setItemPosition(100,10,True)
        composerLabel.adjustSizeToText()
        composerLabel.setMargin(-6)
        c.addItem(composerLabel)

        #scale bar
        item = QgsComposerScaleBar(c)
        item.setStyle('Single Box')
        item.setComposerMap(composerMap)
        item.applyDefaultSize()
        item.setAlignment(QgsComposerScaleBar.Left)
        item.setNumSegmentsLeft(0)
        item.setNumSegments(6)
        item.setItemPosition(159, 198)
        c.addItem(item)

        


        printer = QPrinter()
        printer.setOutputFormat(QPrinter.PdfFormat)
        printer.setOutputFileName('/home/khalifi/Desktop/Resultat/carte.pdf')
        printer.setPaperSize(QSizeF(c.paperWidth(), c.paperHeight()), QPrinter.Millimeter)
        printer.setFullPage(True)
        printer.setColorMode(QPrinter.Color)
        printer.setResolution(c.printResolution())
        pdfPainter = QPainter(printer)
        paperRectMM = printer.pageRect(QPrinter.Millimeter)
        paperRectPixel = printer.pageRect(QPrinter.DevicePixel)
        c.render(pdfPainter, paperRectPixel, paperRectMM)
        pdfPainter.end()


    def exportToWord(self):

        document = docx.Document()
        title=document.add_heading('Agence Urbaine de Khémisset - Dérogation'.decode('utf-8'), 0)

        p = document.add_paragraph('Ce document synthétise les projets dérogés en intersection avec votre projet '.decode('utf-8'))
        document.add_heading('Le résulat'.decode('utf-8'), level=1)
        h= document.add_paragraph("LE NOMBRE DES ZONES DROGEES TOUCHANT LA COUCHE "+self.dlg.comboBoxLayers.currentText()+" EST : " + str(self.dlg.tableWidget.rowCount()) + " .")
        document.add_paragraph("Vous trouverez, ci-dessous, les propriétés de chaque zones :".decode('utf-8'))

        l = self.dlg.tableWidget.rowCount()
        c = self.dlg.tableWidget.columnCount()
        cprjdrg = ["OBJECTID", "REGIME_FON", "REFERENCE_", "CERCLE", "COMMUNE", "STATUT_FON", "SUPERFICIE"]

        def lireTabDerog(l, c, cprjdrg):
            data = []
            donnees = []
        
            for il in range(l):
                for ic in range(c):
                    if self.dlg.tableWidget.horizontalHeaderItem(ic).text() in cprjdrg:
                        data.append(self.dlg.tableWidget.item( il, ic).text())
                donnees.append(data)
                data = []
            return donnees

        table = document.add_table(rows=1, cols=7)

        heading_cells = table.rows[0].cells
        for i in range(len(cprjdrg)):
            heading_cells[i].text = str(cprjdrg[i])
       

        prjsdrj = lireTabDerog(l, c, cprjdrg)
        for d in range(len(prjsdrj)):
            cells = table.add_row().cells
            i=0
            for i in range(7):
                cells[i].text = str(prjsdrj[d][i])

        document.save('/home/khalifi/Desktop/Resultat/resultat.docx')
        


           
    def AfficherIntersection(self):
        if self.h==1:
            

            canvas = self.iface.mapCanvas()
            areas = []
            
            line_layer = self.FindLayerByName('buffer')
            area_layer = self.FindLayerByName(self.dlg.comboBoxLayers.currentText())
            prov = area_layer.dataProvider()
            fields = prov.fields()
            self.dlg.tableWidget.setColumnCount(7)
            self.dlg.tableWidget_2.setColumnCount(1)
            self.dlg.tableWidget.setRowCount(1)
            self.dlg.tableWidget_2.setRowCount(1)

            
            i=0
            for field in fields:
                
                if field.name() == "OBJECTID" or field.name() == "REGIME_FON" or field.name() == "REFERENCE_" or field.name() == "CERCLE" or field.name() == "COMMUNE" or field.name() == "STATUT_FON" or field.name() == "SUPERFICIE" :
                    self.dlg.tableWidget.setHorizontalHeaderItem (i,QTableWidgetItem(field.name()))
                    i=i+1 
               
                
            i=0
            self.dlg.tableWidget_2.setHorizontalHeaderItem (i,QTableWidgetItem("Proportion"))
            for line_feature in line_layer.getFeatures():
                i=0
                for area_feature in area_layer.getFeatures():
                   
                    if (area_feature.geometry().intersects(line_feature.geometry())==True) :
                        areas.append(area_feature.id())
                        self.dlg.tableWidget.insertRow(i+1)
                        self.dlg.tableWidget_2.insertRow(i+1)
                        #self.dlg.tableW.setRowCount(j)
                        geom=line_feature.geometry().intersection(area_feature.geometry())
                        percent=(float(geom.area())/float(area_feature.geometry().area()))*100
                       
                        self.dlg.tableWidget_2.setItem(i,0,QTableWidgetItem(str(percent)))


                        j=0
                        k=0
                        for field in fields:
                            
                            if field.name() == "OBJECTID" or field.name() == "REGIME_FON" or field.name() == "REFERENCE_" or field.name() == "CERCLE" or field.name() == "COMMUNE" or field.name() == "STATUT_FON" or field.name() == "SUPERFICIE" :
                                s=''
                                try :  
                                    s=str(area_feature.attributes()[j])
                                  
                                    self.dlg.tableWidget.setItem(i,k,QTableWidgetItem(s))
                                except :
                                    pass
                                k=k+1
                            j=j+1


                        
                            #print  str(area_feature.attributes()[m])
                        i=i+1
                        
                self.dlg.tableWidget.removeRow (i)
                self.dlg.tableWidget_2.removeRow (i)
            #self.dlg.tableWidget.setItem(i,m+1,QTableWidgetItem(percent))    
            legend = self.iface.legendInterface() 
           # for layer in canvas.layers():
                #if layer.type() == layer.VectorLayer and layer.name()!="OpenStreetMap":
                    #layer.removeSelection()
                     # access the legend
                    #legend.setLayerVisible(layer, False) 
            canvas.refresh()
            legend.setLayerVisible(area_layer, True)
            legend.setLayerVisible(line_layer, True)
            project=self.FindLayerByName('couche_projet')
            legend.setLayerVisible(project, True)
            area_layer.select(areas)
            canvas.zoomToSelected(area_layer)

            self.dlg.label_3.setText("LE NOMBRE DES ZONES DEROGEES TOUCHANT LA COUCHE "+self.dlg.comboBoxLayers.currentText()+" EST : " + str(self.dlg.tableWidget.rowCount()) + " !")
